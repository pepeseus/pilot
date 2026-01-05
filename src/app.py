import streamlit as st
import pandas as pd
from docx import Document
import json
import re

from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from mapper import normalize_header, map_row_data

st.set_page_config(page_title="Flexible Exports Pilot", layout="wide", initial_sidebar_state="collapsed")


# ============================================================
# WordDocument Class - OO approach to document processing
# ============================================================

class WordDocument:
    """
    Encapsulates a Word document with its tree structure and metadata.
    Provides methods to query and extract information from the document.
    """
    
    def __init__(self, docx_file):
        """Initialize with a python-docx Document object."""
        self.doc = Document(docx_file)
        self.tree = self._walk_container(self.doc, "doc")
        self.structure = self._extract_structure()
        self.text_nodes = self._collect_text_nodes()
        self.tables = [x for x in self.structure if x["type"] == "table"]
        self.table_section_map = {f"doc/table[{i}]": t["section"] for i, t in enumerate(self.tables)}
    
    def _iter_block_items(self):
        """Yield paragraphs and tables in document order."""
        for child in self.doc.element.body:
            if isinstance(child, CT_P):
                yield Paragraph(child, self.doc)
            elif isinstance(child, CT_Tbl):
                yield Table(child, self.doc)
    
    def _walk_container(self, container, path):
        """Recursively walk a container (doc or cell) to build a tree of nodes."""
        nodes = []
        
        # Paragraphs in this container
        for i, p in enumerate(container.paragraphs):
            if p.text.strip():
                nodes.append({
                    "type": "paragraph",
                    "text": p.text,
                    "style": p.style.name if p.style else None,
                    "path": f"{path}/p[{i}]",
                    "obj": p
                })
        
        # Tables in this container
        for ti, table in enumerate(container.tables):
            table_path = f"{path}/table[{ti}]"
            table_node = {
                "type": "table",
                "path": table_path,
                "rows": []
            }
            
            for ri, row in enumerate(table.rows):
                row_node = {"type": "row", "cells": []}
                
                for ci, cell in enumerate(row.cells):
                    cell_path = f"{table_path}/row[{ri}]/cell[{ci}]"
                    cell_node = {
                        "type": "cell",
                        "path": cell_path,
                        "children": self._walk_container(cell, cell_path)
                    }
                    row_node["cells"].append(cell_node)
                
                table_node["rows"].append(row_node)
            
            nodes.append(table_node)
        
        return nodes
    
    def _extract_structure(self):
        """Extract high-level document structure with section tracking."""
        structure = []
        current_section = None
        
        for block in self._iter_block_items():
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not text:
                    continue
                
                style = block.style.name if block.style else ""
                
                is_heading = False
                level = None
                
                if style.lower().startswith("heading"):
                    is_heading = True
                    m = re.search(r"(\d+)", style)
                    if m:
                        level = int(m.group(1))
                
                m2 = re.search(r"section\s*(\d+)", text, re.I)
                if m2:
                    is_heading = True
                    current_section = f"section_{m2.group(1).zfill(2)}"
                
                structure.append({
                    "type": "heading" if is_heading else "paragraph",
                    "text": text,
                    "heading_level": level,
                    "style": style,
                    "section": current_section
                })
            
            elif isinstance(block, Table):
                headers = [cell.text.strip() for cell in block.rows[0].cells if cell.text.strip()]
                structure.append({
                    "type": "table",
                    "table": block,
                    "rows": len(block.rows),
                    "cols": len(block.columns),
                    "headers": headers,
                    "section": current_section
                })
        
        return structure
    
    def _collect_text_nodes(self):
        """Collect text-bearing items (headings, paragraphs, table headers) for mapping."""
        items = []
        for idx, item in enumerate(self.structure):
            section = item.get("section")
            if item["type"] in ("heading", "paragraph"):
                text = item["text"].strip()
                if not text:
                    continue
                label = f"{text[:60]}{'...' if len(text) > 60 else ''}"
                items.append({
                    "label": label,
                    "section": section,
                    "text": text,
                    "path": f"item[{idx}]"
                })
            elif item["type"] == "table" and item.get("headers"):
                for h_idx, header in enumerate(item["headers"]):
                    label = f"{header} (Table header)"
                    items.append({
                        "label": label,
                        "section": section,
                        "text": header,
                        "path": f"table[{idx}]/header[{h_idx}]"
                    })
        return items
    
    def get_text_options(self, section=None):
        """Get all text options, optionally filtered by section."""
        if section is None:
            return self.text_nodes
        return [n for n in self.text_nodes if n["section"] == section]
    
    def get_sections(self):
        """Return all unique section identifiers found in the document."""
        return sorted(set(item.get("section") for item in self.structure if item.get("section")))
    
    def find_node_by_path(self, path):
        """Find a node in the tree by its path string."""
        # Simple recursive search implementation
        def search(nodes, target_path):
            for node in nodes:
                if node.get("path") == target_path:
                    return node
                if node.get("type") == "table":
                    for row in node.get("rows", []):
                        for cell in row.get("cells", []):
                            result = search(cell.get("children", []), target_path)
                            if result:
                                return result
            return None
        return search(self.tree, path)
    
    def render_tree_node(self, node):
        """Render a tree node in Streamlit (for debug/visualization)."""
        if node["type"] == "paragraph":
            st.markdown(f"**{node['path']}**")
            st.text(node["text"])
        elif node["type"] == "table":
            with st.expander(f"📊 {node['path']}"):
                for row in node["rows"]:
                    for cell in row["cells"]:
                        with st.expander(cell["path"]):
                            for child in cell["children"]:
                                self.render_tree_node(child)
    
    def extract_to_json(self, mapping):
        """
        Extract data from the Word document into JSON format using the field mappings.
        
        Args:
            mapping: dict from JSON path to text node info
                     e.g., {"section_01.title": {"path": "item[3]", "text": "Title", ...}}
        
        Returns:
            tuple of (json_data dict, results dict with success/errors)
        """
        json_output = {}
        results = {"success": [], "errors": []}
        
        for json_path, word_location in mapping.items():
            try:
                # Extract text from Word document at the mapped location
                text_value = self._extract_text_at_path(word_location["path"])
                
                if text_value is None:
                    results["errors"].append(f"Could not extract text from {word_location['path']}")
                    continue
                
                # Set the value in the nested JSON structure
                self._set_nested_value(json_output, json_path, text_value)
                results["success"].append(f"Extracted {json_path} ← {word_location['path']}")
            except Exception as e:
                results["errors"].append(f"Error extracting {json_path}: {str(e)}")
        
        return json_output, results
    
    def _set_nested_value(self, data, path, value):
        """Set a value in nested dict using dot notation path."""
        path_clean = path.replace("[]", "")
        parts = path_clean.split(".")
        current = data
        
        for i, part in enumerate(parts[:-1]):
            if part not in current:
                current[part] = {}
            current = current[part]
        
        # Set the final value
        current[parts[-1]] = value
    
    def _extract_text_at_path(self, path):
        """Extract text from a specific path in the document."""
        try:
            if path.startswith("item["):
                # It's a paragraph/heading in the structure
                idx = int(path.split("[")[1].split("]")[0])
                if idx >= len(self.structure):
                    return None
                
                item = self.structure[idx]
                if item["type"] in ("heading", "paragraph"):
                    return item["text"]
            
            elif "table" in path and "header" in path:
                # It's a table header cell
                table_idx = int(path.split("table[")[1].split("]")[0])
                header_idx = int(path.split("header[")[1].split("]")[0])
                
                if table_idx >= len(self.tables):
                    return None
                
                table_item = self.tables[table_idx]
                table = table_item["table"]
                
                if not table.rows or header_idx >= len(table.rows[0].cells):
                    return None
                
                # Extract from the header cell
                cell = table.rows[0].cells[header_idx]
                if cell.paragraphs:
                    return cell.paragraphs[0].text
        except Exception:
            return None
        
        return None
    
    def _modify_text_at_path(self, path, new_text):
        """Modify text at a specific path in the document."""
        # Parse path like "item[3]" or "table[2]/header[0]"
        if path.startswith("item["):
            # It's a paragraph/heading in the structure
            idx = int(path.split("[")[1].split("]")[0])
            if idx >= len(self.structure):
                raise ValueError(f"Invalid item index: {idx}")
            
            item = self.structure[idx]
            if item["type"] in ("heading", "paragraph"):
                # Find the actual paragraph object in the tree
                tree_path = self._structure_to_tree_path(idx)
                node = self.find_node_by_path(tree_path)
                if node and "obj" in node:
                    para = node["obj"]
                    # Preserve formatting by keeping the style and font properties
                    if para.runs:
                        # Get formatting from first run
                        first_run = para.runs[0]
                        bold = first_run.bold
                        italic = first_run.italic
                        font_name = first_run.font.name
                        font_size = first_run.font.size
                        
                        # Clear and re-add with formatting
                        para.clear()
                        run = para.add_run(new_text)
                        run.bold = bold
                        run.italic = italic
                        if font_name:
                            run.font.name = font_name
                        if font_size:
                            run.font.size = font_size
                    else:
                        para.text = new_text
        
        elif "table" in path and "header" in path:
            # It's a table header cell
            # Parse: "table[2]/header[0]"
            table_idx = int(path.split("table[")[1].split("]")[0])
            header_idx = int(path.split("header[")[1].split("]")[0])
            
            if table_idx >= len(self.tables):
                raise ValueError(f"Invalid table index: {table_idx}")
            
            table_item = self.tables[table_idx]
            table = table_item["table"]
            
            if not table.rows or header_idx >= len(table.rows[0].cells):
                raise ValueError(f"Invalid header index: {header_idx}")
            
            # Modify the header cell
            cell = table.rows[0].cells[header_idx]
            if cell.paragraphs:
                para = cell.paragraphs[0]
                if para.runs:
                    # Preserve formatting
                    first_run = para.runs[0]
                    bold = first_run.bold
                    italic = first_run.italic
                    font_name = first_run.font.name
                    font_size = first_run.font.size
                    
                    para.clear()
                    run = para.add_run(new_text)
                    run.bold = bold
                    run.italic = italic
                    if font_name:
                        run.font.name = font_name
                    if font_size:
                        run.font.size = font_size
                else:
                    para.text = new_text
    
    def _structure_to_tree_path(self, struct_idx):
        """Convert structure index to tree path."""
        # Count paragraphs before this structure item
        para_count = 0
        for i, item in enumerate(self.structure[:struct_idx + 1]):
            if item["type"] in ("heading", "paragraph"):
                if i == struct_idx:
                    return f"doc/p[{para_count}]"
                para_count += 1
        return None
    
    def save(self, output_path):
        """Save the modified document to a file path or BytesIO buffer."""
        self.doc.save(output_path)

# ============================================================
# JSON Schema logic (unchanged)
# ============================================================

def resolve_schema_ref(root_schema, ref_path):
    if not ref_path.startswith("#/"):
        return None
    parts = ref_path[2:].split("/")
    current = root_schema
    for part in parts:
        if not isinstance(current, dict):
            return None
        current = current.get(part)
        if current is None:
            return None
    return current

def extract_json_paths(schema_data, prefix="", section="", root=None, group=None, seen=None):
    """
    Extract REQUIRED fields from the JSON schema with proper grouping.
    Only properties listed in their parent's `required` array are returned.
    """
    if root is None:
        root = schema_data
    if not isinstance(schema_data, dict):
        return []
    if seen is None:
        seen = set()

    paths = []

    def add_leaf(path, field_name, current_section, type_name):
        top_group = group or (path.split(".")[0] if path else None)
        # avoid duplicates
        if path in seen:
            return
        seen.add(path)
        paths.append({
            "path": path,
            "field_name": field_name,
            "section": current_section,
            "type": type_name or "unknown",
            "group": top_group
        })

    required_set = set(schema_data.get("required", [])) if isinstance(schema_data.get("required", []), list) else set()
    # If no required specified, include all properties at this level (so we don't drop fields)
    include_all_if_no_required = len(required_set) == 0

    if "properties" in schema_data:
        for k, v in schema_data["properties"].items():
            # Skip non-required properties only when required list is present
            if not include_all_if_no_required and k not in required_set:
                continue

            path = f"{prefix}.{k}" if prefix else k
            current_section = k if k.startswith("section_") else section
            current_group = group or k

            # Handle anyOf/allOf/oneOf with $ref inside
            combo = v.get("anyOf") or v.get("allOf") or v.get("oneOf")
            if combo:
                for option in combo:
                    if option.get("type") == "null":
                        continue  # skip null alternative
                    if "$ref" in option:
                        ref = resolve_schema_ref(root, option["$ref"])
                        if ref:
                            paths += extract_json_paths(ref, path, current_section, root, current_group)
                    elif option.get("type"):
                        # Treat as leaf if simple
                        if option.get("type") != "object" or "properties" not in option:
                            add_leaf(path, k, current_section, option.get("type"))
                        else:
                            paths += extract_json_paths(option, path, current_section, root, current_group)
                continue

            if "$ref" in v:
                ref = resolve_schema_ref(root, v["$ref"])
                if ref:
                    paths += extract_json_paths(ref, path, current_section, root, current_group)
            elif v.get("type") == "object":
                paths += extract_json_paths(v, path, current_section, root, current_group)
            elif v.get("type") == "array":
                items = v.get("items")
                # Record the array itself
                add_leaf(path, k, current_section, "array")
                if isinstance(items, dict) and "$ref" in items:
                    ref = resolve_schema_ref(root, items["$ref"])
                    if ref:
                        paths += extract_json_paths(ref, f"{path}[]", current_section, root, current_group)
                elif isinstance(items, dict):
                    if items.get("type") and (items.get("type") != "object" or "properties" not in items):
                        add_leaf(f"{path}[]", k, current_section, items.get("type"))
                    else:
                        paths += extract_json_paths(items, f"{path}[]", current_section, root, current_group)
            else:
                # Leaf node (any type)
                add_leaf(path, k, current_section, v.get("type"))

    return paths

# ============================================================
# Streamlit App
# ============================================================

st.title("Flexible Exports Pilot")

template_file = st.file_uploader("Upload Word Template", type="docx")
schema_file = st.file_uploader("Upload JSON Schema", type="json")

if template_file and schema_file:
    schema_data = json.load(schema_file)
    
    # Initialize WordDocument (parses everything once)
    word_doc = WordDocument(template_file)
    
    st.success(f"✓ Found {len(word_doc.tables)} table(s)")

    with st.expander("📄 View Document Structure"):
        for item in word_doc.structure:
            if item["type"] == "heading":
                st.markdown(f"### 📌 {item['text']}")
            elif item["type"] == "paragraph":
                st.text(item["text"])
            else:
                st.markdown(f"📊 Table ({item['rows']}×{item['cols']}) — Section: {item['section']}")
                if item["headers"]:
                    st.caption(", ".join(item["headers"]))

    # Full DOM walk (paragraphs, tables, cells)
    with st.expander("📄 Full Word DOM", expanded=False):
        st.caption("Nested view of all paragraphs, tables, rows, and cells with paths")
        for n in word_doc.tree:
            word_doc.render_tree_node(n)

    # Extract Word columns
    word_columns = []
    for idx, t in enumerate(word_doc.tables):
        for cell in t["table"].rows[0].cells:
            if cell.text.strip():
                word_columns.append({
                    "column_name": cell.text.strip(),
                    "table_index": idx,
                    "section": t["section"]
                })

    json_paths = extract_json_paths(schema_data)

    word_col_display = []
    word_lookup = {}
    for c in word_columns:
        label = f"{c['column_name']} ({c['section'] or 'No section'})"
        word_col_display.append(label)
        word_lookup[label] = {
            **c,
            "type": "header_cell",
            "path": f"table[{c['table_index']}]/header/{c['column_name']}"
        }

    # Text options (paragraphs, headings, table headers) by section
    text_by_section = {}
    for n in word_doc.text_nodes:
        text_by_section.setdefault(n["section"], []).append(n)

    st.subheader("Review Mappings (required fields → Word text)")

    selections = {}

    # Group fields by top-level key (supports arbitrary schemas, not just section_*)
    fields_by_group = {}
    for j in json_paths:
        grp = j.get("group") or "ungrouped"
        fields_by_group.setdefault(grp, []).append(j)

    # Options: all text nodes, no filtering (bare bones)
    all_option_labels = [None] + [o["label"] for o in word_doc.text_nodes]
    all_option_lookup = {o["label"]: o for o in word_doc.text_nodes}

    for group_name, fields in fields_by_group.items():
        st.markdown(f"#### Group: {group_name}")

        data = []
        for j in fields:
            # auto-guess exact text match across all options
            guess_label = None
            for o in word_doc.text_nodes:
                if normalize_header(o["text"]) == normalize_header(j["field_name"]):
                    guess_label = o["label"]
                    break
            data.append({
                "Field": j["field_name"],
                "_json_path": j["path"],  # hidden helper column
                "Word Text": guess_label
            })

        df = pd.DataFrame(data)

        edited = st.data_editor(
            df,
            column_config={
                "Word Text": st.column_config.SelectboxColumn(
                    options=all_option_labels,
                    help="Select any text from the Word document"
                ),
                "Field": st.column_config.TextColumn(disabled=True),
                "_json_path": None  # hide helper column
            },
            hide_index=True
        )

        # Collect selections
        for _, row in edited.iterrows():
            if pd.notna(row["Word Text"]):
                selections[row["_json_path"]] = {
                    "field": row["Field"],
                    **(all_option_lookup.get(row["Word Text"], {}))
                }

    st.markdown("---")
    st.subheader("Extract Data to JSON")
    
    if st.button("Extract Data from Word Document", type="primary", disabled=not selections):
        if selections:
            # Extract data from Word document
            with st.spinner("Extracting data from document..."):
                json_output, results = word_doc.extract_to_json(selections)
            
            # Show results
            if results["success"]:
                st.success(f"✓ Successfully extracted {len(results['success'])} field(s)")
                with st.expander("View successful extractions"):
                    for msg in results["success"]:
                        st.text(f"✓ {msg}")
            
            if results["errors"]:
                st.warning(f"⚠️ {len(results['errors'])} error(s) occurred")
                with st.expander("View errors"):
                    for msg in results["errors"]:
                        st.text(f"✗ {msg}")
            
            # Show extracted JSON
            st.subheader("Extracted JSON Data")
            st.json(json_output)
            
            # Create download button for JSON
            json_str = json.dumps(json_output, indent=2)
            st.download_button(
                label="📥 Download JSON Data",
                data=json_str,
                file_name="extracted_data.json",
                mime="application/json",
                key="download_json"
            )
        else:
            st.warning("⚠️ No field mappings configured. Please map fields above first.")
        
    # Debug: Show current mappings
    with st.expander("🔍 Debug: View Current Mappings"):
        st.json(selections)
