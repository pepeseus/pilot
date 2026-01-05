import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import json
from datetime import datetime
from io import BytesIO
import re
import html

st.set_page_config(page_title="Document Generator", layout="wide")

# ============================================================
# HTML TO WORD FORMATTING HELPER
# ============================================================

def apply_html_to_paragraph(paragraph, html_text):
    """
    Parse HTML and apply formatting to Word paragraph.
    Supports: <b>, <strong>, <i>, <em>, <u>, <br>

    To support rich text (HTML)!
    """
    paragraph.clear()
    
    # Track current formatting state
    is_bold = False
    is_italic = False
    is_underline = False
    
    # Pattern to match HTML tags
    tag_pattern = r'<(/?)([a-z]+)>'
    
    last_end = 0
    for match in re.finditer(tag_pattern, html_text, re.IGNORECASE):
        # Add text before this tag
        text_before = html_text[last_end:match.start()]
        if text_before:
            text_before = html.unescape(text_before)
            run = paragraph.add_run(text_before)
            run.bold = is_bold
            run.italic = is_italic
            run.underline = is_underline
        
        # Process the tag
        is_closing = match.group(1) == '/'
        tag_name = match.group(2).lower()
        
        if tag_name == 'br':
            paragraph.add_run('\n')
        elif tag_name in ['b', 'strong']:
            is_bold = not is_closing
        elif tag_name in ['i', 'em']:
            is_italic = not is_closing
        elif tag_name == 'u':
            is_underline = not is_closing
        
        last_end = match.end()
    
    # add remaining text after last tag
    text_after = html_text[last_end:]
    if text_after:
        text_after = html.unescape(text_after)
        run = paragraph.add_run(text_after)
        run.bold = is_bold
        run.italic = is_italic
        run.underline = is_underline
    
    # if no runs were added, add the original text as-is
    if not paragraph.runs:
        paragraph.add_run(html.unescape(html_text))

def has_html_tags(text):
    """Check if text contains HTML tags."""
    return bool(re.search(r'<[^>]+>', text))

st.title("📄 Document Generator")
st.markdown("Upload your template, mapping config, and optionally existing data - then fill in or edit the values to generate your document.")

# ============================================================
# FILE UPLOADERS
# ============================================================

col1, col2, col3 = st.columns(3)

with col1:
    template_file = st.file_uploader("📄 Word Template", type="docx", help="The Word document template")

with col2:
    mapping_file = st.file_uploader("🗺️ Mapping Config", type="json", help="The mapping configuration from the mapper")

with col3:
    data_file = st.file_uploader("📋 Existing Data (Optional)", type="json", help="Pre-fill with existing JSON data")

if template_file and mapping_file:
    # Load files
    doc = Document(template_file)
    mapping_config = json.load(mapping_file)
    
    # Load existing data if provided
    existing_data = {}
    data_file_name = None
    if data_file:
        existing_data = json.load(data_file)
        data_file_name = data_file.name
        st.success(f"✓ Loaded existing data with {len(existing_data)} top-level keys")
    
    st.markdown("---")
    
    # Helper function to extract value from nested JSON with array support
    def extract_value_from_data(json_path, data):
        """Extract value from nested data structure, handling arrays."""
        # Handle array notation
        has_array = "[]" in json_path
        path_parts = json_path.replace("[]", "").split(".")
        
        current = data
        for i, part in enumerate(path_parts):
            if not isinstance(current, dict):
                # If we hit an array when we expected a dict
                if isinstance(current, list) and len(current) > 0:
                    # For array fields, take the first item as default
                    current = current[0]
                    if isinstance(current, dict):
                        current = current.get(part, "")
                    else:
                        return ""
                else:
                    return ""
            else:
                current = current.get(part, "")
                
                # If this is an array and we're not at the last part
                if isinstance(current, list) and i < len(path_parts) - 1:
                    if len(current) > 0:
                        # Take first item for now
                        current = current[0]
                    else:
                        return ""
        
        return current if current else ""
    
    # Extract fields from mapping config
    fields_data = []
    for json_path, config in mapping_config.items():
        field_name = config["field_name"]
        field_format = config.get("format")
        group = config.get("group", "ungrouped")
        
        # Get existing value if available
        existing_value = extract_value_from_data(json_path, existing_data)
        
        if not existing_value:
            existing_value = ""
        
        fields_data.append({
            "path": json_path,
            "name": field_name,
            "format": field_format,
            "group": group,
            "value": str(existing_value),
            "doc_location": config.get("document_location", {})
        })
    
    # Group fields by section
    fields_by_group = {}
    for field in fields_data:
        grp = field["group"]
        fields_by_group.setdefault(grp, []).append(field)
    
    # Display form for each group
    st.subheader("📝 Enter/Edit Data")
    st.caption("Fill in the values for each field. Date fields show a calendar picker 📅, email fields validate format 📧")
    
    with st.expander("💡 Rich Text Formatting (HTML)"):
        st.markdown("""
        **Text fields support HTML formatting:**
        - `<b>bold text</b>` or `<strong>bold text</strong>` → **bold text**
        - `<i>italic text</i>` or `<em>italic text</em>` → *italic text*
        - `<u>underlined text</u>` → <u>underlined text</u>
        - `<br>` → line break
        
        **Example:**
        ```
        This is <b>bold</b>, this is <i>italic</i>, and this is <u>underlined</u>.<br>
        This is a new line with <b><i>bold italic</i></b> text.
        ```
        """)
    
    # Store values in session state - update if data file changes
    if "last_data_file" not in st.session_state:
        st.session_state.last_data_file = None
    
    # Check if we need to reload data (new file or first time)
    data_changed = st.session_state.last_data_file != data_file_name
    
    if "field_values" not in st.session_state:
        st.session_state.field_values = {}
    
    # Update field values and widget keys when data changes
    if data_changed:
        st.session_state.last_data_file = data_file_name
        
        # Update our tracking dict AND the widget state keys
        for field in fields_data:
            path = field["path"]
            value = field["value"]
            field_format = field.get("format")
            
            # Store in our tracking
            st.session_state.field_values[path] = value
            
            # Also set the widget keys directly
            if field_format == "date":
                # For dates, parse and store as date object
                if value:
                    try:
                        st.session_state[f"date_{path}"] = datetime.strptime(value, "%Y-%m-%d").date()
                    except:
                        st.session_state[f"date_{path}"] = datetime.now().date()
            elif field_format == "email":
                st.session_state[f"email_{path}"] = value
            else:
                st.session_state[f"text_{path}"] = value
        
        if data_file_name:
            filled_count = len([f for f in fields_data if f['value']])
            st.info(f"📥 Pre-filled {filled_count} / {len(fields_data)} fields from {data_file_name}")
            
            # Show which fields were filled
            with st.expander("🔍 View Pre-filled Fields"):
                for field in fields_data:
                    if field['value']:
                        value_preview = field['value'][:50] + "..." if len(field['value']) > 50 else field['value']
                        st.text(f"✓ {field['path']}: {value_preview}")
    
    # Create tabs for each group
    group_tabs = st.tabs([f"📁 {grp}" for grp in sorted(fields_by_group.keys())])
    
    for tab, (group_name, fields) in zip(group_tabs, sorted(fields_by_group.items())):
        with tab:
            for field in fields:
                path = field["path"]
                name = field["name"]
                field_format = field["format"]
                current_value = st.session_state.field_values.get(path, "")
                
                # Create input based on format
                if field_format == "date":
                    # Date picker
                    label = f"📅 {path}"
                    try:
                        # Try to parse existing date
                        if current_value:
                            default_date = datetime.strptime(current_value, "%Y-%m-%d").date()
                        else:
                            default_date = datetime.now().date()
                    except:
                        default_date = datetime.now().date()
                    
                    date_value = st.date_input(
                        label,
                        value=default_date,
                        key=f"date_{path}"
                    )
                    st.session_state.field_values[path] = date_value.strftime("%Y-%m-%d")
                
                elif field_format == "email":
                    # Email input with validation
                    label = f"📧 {path}"
                    email_value = st.text_input(
                        label,
                        value=current_value,
                        key=f"email_{path}",
                        placeholder="user@example.com"
                    )
                    
                    # Validate email
                    if email_value:
                        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                        if not re.match(email_pattern, email_value):
                            st.warning("⚠️ Invalid email format")
                    
                    st.session_state.field_values[path] = email_value
                
                else:
                    # Regular text input - use text_area for HTML support
                    text_value = st.text_area(
                        path,
                        value=current_value,
                        key=f"text_{path}",
                        height=100,
                        help="Supports HTML: <b>bold</b>, <i>italic</i>, <u>underline</u>, <br> for line break"
                    )
                    st.session_state.field_values[path] = text_value
                    
                    # Show HTML preview if HTML tags detected
                    if text_value and has_html_tags(text_value):
                        st.caption("🎨 HTML detected - will apply formatting")
                        # Show a preview of how it will render
                        preview_html = text_value.replace('<br>', '<br/>')
                        st.markdown(f"Preview: {preview_html}", unsafe_allow_html=True)
    
    # Show data preview
    st.markdown("---")
    
    with st.expander("🔍 Preview JSON Data"):
        # Build the nested JSON structure
        output_json = {}
        for path, value in st.session_state.field_values.items():
            if value:  # Only include non-empty values
                parts = path.replace("[]", "").split(".")
                current = output_json
                for part in parts[:-1]:
                    if part not in current:
                        current[part] = {}
                    current = current[part]
                current[parts[-1]] = value
        
        st.json(output_json)
    
    # Generate document
    st.markdown("---")
    col1, col2 = st.columns([1, 3])
    
    with col1:
        if st.button("🚀 Generate Document", type="primary"):
            # Build the nested JSON
            output_json = {}
            for path, value in st.session_state.field_values.items():
                if value:
                    parts = path.replace("[]", "").split(".")
                    current = output_json
                    for part in parts[:-1]:
                        if part not in current:
                            current[part] = {}
                        current = current[part]
                    current[parts[-1]] = value
            
            # Apply values to document
            # Parse document to get all text segments (reusing logic from mapper)
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
            
            segments = []
            for child in doc.element.body:
                if isinstance(child, CT_P):
                    para = Paragraph(child, doc)
                    if para.text.strip():
                        segments.append({"type": "paragraph", "obj": para})
                elif isinstance(child, CT_Tbl):
                    table = Table(child, doc)
                    seen_in_row = {}
                    for row_idx, row in enumerate(table.rows):
                        if row_idx not in seen_in_row:
                            seen_in_row[row_idx] = set()
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text and text not in seen_in_row[row_idx]:
                                seen_in_row[row_idx].add(text)
                                if cell.paragraphs:
                                    segments.append({"type": "cell", "obj": cell.paragraphs[0]})
            
            # Apply mappings
            updated_count = 0
            debug_info = []
            
            for path, value in st.session_state.field_values.items():
                if value and path in mapping_config:
                    doc_loc = mapping_config[path]["document_location"]
                    idx = doc_loc["index"]
                    
                    if idx < len(segments):
                        para = segments[idx]["obj"]
                        
                        # Preserve original formatting attributes
                        original_font_name = None
                        original_font_size = None
                        if para.runs:
                            first_run = para.runs[0]
                            original_font_name = first_run.font.name
                            original_font_size = first_run.font.size
                        
                        # Check if value contains HTML tags
                        value_str = str(value)
                        has_html = has_html_tags(value_str)
                        
                        if has_html:
                            # Apply HTML formatting
                            debug_info.append(f"🎨 Applied HTML to {path}: {value_str[:50]}...")
                            apply_html_to_paragraph(para, value_str)
                            
                            # Reapply original font properties to all runs
                            if original_font_name or original_font_size:
                                for run in para.runs:
                                    if original_font_name:
                                        run.font.name = original_font_name
                                    if original_font_size:
                                        run.font.size = original_font_size
                        else:
                            # Plain text - preserve formatting
                            debug_info.append(f"📝 Applied plain text to {path}: {value_str[:50]}...")
                            if para.runs:
                                first_run = para.runs[0]
                                bold = first_run.bold
                                italic = first_run.italic
                                
                                para.clear()
                                run = para.add_run(value_str)
                                run.bold = bold
                                run.italic = italic
                                if original_font_name:
                                    run.font.name = original_font_name
                                if original_font_size:
                                    run.font.size = original_font_size
                            else:
                                para.text = value_str
                        
                        updated_count += 1
            
            st.success(f"✅ Generated document! Updated {updated_count} fields.")
            
            # Show debug info
            with st.expander("🔍 Debug Info - What was applied"):
                for info in debug_info:
                    st.text(info)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Download buttons
            col_a, col_b = st.columns(2)
            
            with col_a:
                st.download_button(
                    "📥 Download Word Document",
                    data=buffer,
                    file_name="generated_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col_b:
                json_str = json.dumps(output_json, indent=2)
                st.download_button(
                    "📥 Download JSON Data",
                    data=json_str,
                    file_name="document_data.json",
                    mime="application/json"
                )
    
    with col2:
        st.info("👈 Click to generate the Word document with your data")

else:
    st.info("👆 Upload a Word template and mapping configuration to begin")
    
    with st.expander("ℹ️ How to use"):
        st.markdown("""
        ### Step-by-step guide:
        
        1. **Upload Word Template**: The `.docx` file you want to populate
        2. **Upload Mapping Config**: The `.json` file from the Interactive Mapper
        3. **Upload Existing Data (Optional)**: Pre-fill fields with a `.json` data file
        
        ### Features:
        - 📅 **Date fields**: Use calendar picker for easy date selection
        - 📧 **Email fields**: Automatic format validation
        - 📁 **Grouped by section**: Fields organized by schema groups
        - 👁️ **Preview JSON**: See the complete JSON structure before generating
        - 💾 **Download both**: Get the Word document AND the JSON data
        
        ### Workflow:
        1. Fill in all required fields (or edit existing values)
        2. Preview the JSON structure
        3. Click "Generate Document"
        4. Download the populated Word document
        5. Optionally download the JSON for future use
        """)

